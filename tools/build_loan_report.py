from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documentation"
TMP_DIR = ROOT / "tmp" / "documentation"
OUT_PATH = OUT_DIR / "Loan_Approval_Prediction_Assignment_Documentation.docx"
WORKFLOW_PATH = TMP_DIR / "loan_workflow.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
GRAY = "555555"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "000000"
GOLD = "7A5A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char, instr_text, fld_sep, text, fld_end))
    return run


def add_bottom_border(paragraph, color="D9E2F3", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc):
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True

    configure_header_footer(section)


def configure_header_footer(section):
    for header in (section.header, section.even_page_header):
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("BMCS2203 Artificial Intelligence")
        set_run_font(r, size=9, color=GRAY, bold=True)
        r = p.add_run("    |    Loan Approval Prediction")
        set_run_font(r, size=9, color=GRAY)
        add_bottom_border(p, color="D9E2F3", size="6")

    for footer in (section.footer, section.even_page_footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(3)
        r = p.add_run("Page ")
        set_run_font(r, size=9, color=GRAY)
        r = add_field(p, "PAGE")
        set_run_font(r, size=9, color=GRAY)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("BMCS2203 ARTIFICIAL INTELLIGENCE"), size=12, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("Loan Approval Prediction"), size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_run_font(
        p.add_run("Using Supervised Machine Learning: A Comparison of Decision Tree and Random Forest"),
        size=14,
        color=DARK_BLUE,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(42)
    set_run_font(p.add_run("Assignment Documentation"), size=11, color=GRAY, italic=True)

    metadata = [
        ("Session", "202605, Year 2026/27"),
        ("Programme", "RSW Y2S1"),
        ("Tutorial Group", "9"),
        ("Tutor", "Mr MOHD HANIF BIN YUSOFF"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for i, (label, value) in enumerate(metadata):
        set_cell_shading(table.cell(i, 0), LIGHT_GRAY)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), size=10.5, color=NAVY, bold=True)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("Team Members"), size=12, color=BLUE, bold=True)
    members = doc.add_table(rows=3, cols=4)
    members.style = "Table Grid"
    set_table_geometry(members, [600, 3000, 3900, 1860])
    headers = ["No.", "Student", "Module in Charge", "Signature / Date"]
    for i, text in enumerate(headers):
        set_cell_shading(members.cell(0, i), BLUE)
        p = members.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=9.5, color=WHITE, bold=True)
    rows = [
        ("1", "Ee Wan Yin\n26WMR12818", "Data preprocessing and Decision Tree", ""),
        ("2", "Lam Chee Weng\n26WMR12849", "Random Forest, evaluation and deployment", ""),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            p = members.cell(r_idx, c_idx).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), size=9.5)

    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), bold=True)
        text = text[len(bold_lead):]
    set_run_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_run_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_run_font(p.add_run(text))
    return p


def add_numbered_list(doc, items):
    numbering = doc.part.numbering_part.element
    decimal_abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        level = abstract.find(qn("w:lvl"))
        if level is None or level.get(qn("w:ilvl")) != "0":
            continue
        num_fmt = level.find(qn("w:numFmt"))
        if num_fmt is not None and num_fmt.get(qn("w:val")) == "decimal":
            decimal_abstract_id = abstract.get(qn("w:abstractNumId"))
            break
    if decimal_abstract_id is None:
        raise ValueError("No decimal numbering definition is available.")

    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(decimal_abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)

    for text in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend((ilvl, num_id_node))
        p_pr.insert(0, num_pr)
        set_run_font(p.add_run(text))


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    set_cell_shading(table.cell(0, 0), CALLOUT)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(f"{label}: "), size=10.5, color=NAVY, bold=True)
    set_run_font(p.add_run(text), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, caption, headers, rows, widths_dxa, aligns=None):
    p = doc.add_paragraph(caption, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), size=9.5, color=NAVY, bold=True)
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            p = table.cell(r_idx, c_idx).paragraphs[0]
            if aligns:
                p.alignment = aligns[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(str(value)), size=9.5)
    if len(rows) <= 4:
        for row in table.rows[:-1]:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def make_workflow():
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1500, 760), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    font_path = Path(r"C:\Windows\Fonts\calibri.ttf")
    bold_path = Path(r"C:\Windows\Fonts\calibrib.ttf")
    font = ImageFont.truetype(str(font_path), 31) if font_path.exists() else ImageFont.load_default()
    bold = ImageFont.truetype(str(bold_path), 34) if bold_path.exists() else font
    small = ImageFont.truetype(str(font_path), 25) if font_path.exists() else font

    boxes = [
        (70, 90, 400, 250, "1. Load and validate", "4,269 applications"),
        (585, 90, 915, 250, "2. Select inputs", "Income, amount, term, CIBIL"),
        (1100, 90, 1430, 250, "3. Stratified split", "70% train / 30% test"),
        (1100, 500, 1430, 660, "4. Model selection", "5-fold GridSearchCV"),
        (585, 500, 915, 660, "5. Final evaluation", "Untouched test set"),
        (70, 500, 400, 660, "6. Save and deploy", "Joblib + Streamlit"),
    ]
    for x1, y1, x2, y2, title, subtitle in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill="#F4F6F9", outline="#2E74B5", width=5)
        title_box = draw.textbbox((0, 0), title, font=bold)
        subtitle_box = draw.textbbox((0, 0), subtitle, font=small)
        draw.text(((x1 + x2 - (title_box[2] - title_box[0])) / 2, y1 + 38), title, font=bold, fill="#203748")
        draw.text(((x1 + x2 - (subtitle_box[2] - subtitle_box[0])) / 2, y1 + 100), subtitle, font=small, fill="#555555")

    arrow_color = "#7A5A00"
    arrows = [((400, 170), (585, 170)), ((915, 170), (1100, 170)), ((1265, 250), (1265, 500)), ((1100, 580), (915, 580)), ((585, 580), (400, 580))]
    for start, end in arrows:
        draw.line((start, end), fill=arrow_color, width=8)
        ex, ey = end
        sx, sy = start
        if ex > sx:
            points = [(ex, ey), (ex - 24, ey - 15), (ex - 24, ey + 15)]
        elif ex < sx:
            points = [(ex, ey), (ex + 24, ey - 15), (ex + 24, ey + 15)]
        else:
            points = [(ex, ey), (ex - 15, ey - 24), (ex + 15, ey - 24)]
        draw.polygon(points, fill=arrow_color)
    image.save(WORKFLOW_PATH, dpi=(180, 180))


def add_intro(doc):
    add_heading(doc, "1. Introduction", 1)
    add_heading(doc, "1.1 Background", 2)
    add_body(doc, "Financial institutions receive many loan applications that must be assessed consistently and efficiently. Each application contains information about the applicant's financial position and the requested loan. When the volume of applications is high, manual checking can be slow, and decisions may vary between evaluators. Supervised machine learning can support this process by learning patterns from previously labelled applications and predicting whether a new application resembles records labelled Approved or Rejected.")
    add_body(doc, "This project develops a binary loan approval prediction prototype using a public dataset of 4,269 applications. Decision Tree and Random Forest classifiers are trained using annual income, loan amount, loan term, and CIBIL score. Both models are evaluated under the same experimental conditions so that their performance can be compared fairly. The final pipelines are saved and integrated into a Streamlit interface that accepts the four required values and displays both predictions.")
    add_callout(doc, "Scope", "The target is the historical loan_status label. It is not a borrower-default or repayment-risk outcome. The system is an educational prototype and must not be used to make real lending decisions.")

    add_heading(doc, "1.2 Problem Statement", 2)
    add_body(doc, "Loan approval decisions depend on relationships among applicant income, requested amount, repayment period, and credit history. Simple single-variable rules can overlook interactions between these values, while manual assessment can be time-consuming and inconsistent. A supervised classifier can identify nonlinear patterns, but different algorithms may produce different levels of accuracy, precision, recall, and F1-score.")
    add_body(doc, "A second problem is methodological. Model results are not comparable when different data splits, preprocessing decisions, or tuning procedures are used. This study therefore trains Decision Tree and Random Forest models using the same four inputs, the same stratified 70:30 split, the same five-fold validation strategy, and the same untouched test set. The models are selected by Approved-class F1-score and reported using only accuracy, precision, recall, and F1-score.")

    add_heading(doc, "1.3 Objectives/Aims", 2)
    objectives = [
        "To prepare and validate the loan approval dataset for supervised binary classification.",
        "To develop Decision Tree and Random Forest pipelines using four required inputs: annual income, loan amount, loan term, and CIBIL score.",
        "To optimize both classifiers using five-fold stratified GridSearchCV and Approved-class F1-score.",
        "To compare the selected models on one untouched test set using accuracy, precision, recall, and F1-score.",
        "To develop a Streamlit application that collects the four inputs and displays the Approved or Rejected prediction from both classifiers.",
    ]
    add_numbered_list(doc, objectives)

    add_heading(doc, "1.4 Significance / Contribution of the Study", 2)
    add_body(doc, "This study demonstrates a complete and reproducible supervised machine-learning workflow for loan approval prediction. Its main methodological contribution is a controlled comparison of Decision Tree and Random Forest classifiers using the same validated dataset, selected input features, stratified data split, cross-validation strategy, and untouched test set.")
    add_body(doc, "The use of scikit-learn Pipelines and GridSearchCV ensures that each model is trained and validated under consistent conditions. Approved-class F1-score is used for model selection because it balances precision and recall, while accuracy, precision, recall, and F1-score are reported for the final evaluation. The reduced four-input design also creates a simple Streamlit interface that collects only the values used by the fitted models.")
    add_body(doc, "Although the prototype is not intended to replace the full assessment procedures used by financial institutions, it provides evidence of how supervised machine learning can support preliminary loan approval prediction. It also establishes a reproducible foundation for future evaluation using better-documented, more diverse, and independently collected lending datasets.")


def add_related_work(doc):
    add_heading(doc, "2. Related Work", 1)
    add_heading(doc, "2.1 Review of Previous Studies", 2)
    add_body(doc, "Research on loan approval prediction commonly compares multiple supervised classifiers because no single method is best for every dataset. Sinap (2024) compared Logistic Regression, K-Nearest Neighbors, Support Vector Machine, Decision Tree, and Random Forest, and reported that Random Forest achieved the strongest result at 97.71% accuracy after feature selection and cross-validation. The study also showed that feature-selection strategy can materially affect measured performance.")
    add_body(doc, "Bhakti, Prasetyo, and Arsi (2024) focused on Random Forest hyperparameter tuning and class balancing for loan approval prediction. Their best experiment achieved 86.2% accuracy. The lower score than some other studies reinforces that results depend on the dataset, preprocessing, class distribution, and evaluation design rather than only on the algorithm name.")
    add_body(doc, "Lakshmi, Gajendra, and Mohanraju (2025) directly compared Decision Tree and Random Forest for loan approval using applicant variables including income, employment, education, CIBIL score, loan amount, and dependents. Their findings favoured Random Forest for accuracy and generalizability. This supports the present project's use of Random Forest as the primary model and Decision Tree as the interpretable comparison model.")
    add_body(doc, "Random Forest was introduced by Breiman (2001) as an ensemble of randomized tree predictors. Combining many trees can reduce the instability of a single Decision Tree and improve generalization when the individual trees are sufficiently strong and not excessively correlated. In contrast, a Decision Tree learns a hierarchy of simple feature-based rules that can be inspected more easily, but an unrestricted tree can overfit its training data. Depth and minimum-leaf constraints are therefore relevant tuning parameters.")

    add_heading(doc, "2.2 Research Gap and Justification for the Current Study", 2)
    add_body(doc, "Previous studies used different datasets, applicant attributes, preprocessing methods, validation procedures, and model configurations. Consequently, differences in their reported results may have been caused by the experimental setup rather than only by the classification algorithms. Some studies also predicted loan approval, while others predicted repayment or default, so their headline scores do not necessarily measure the same task.")
    add_body(doc, "This project addresses these limitations by evaluating Decision Tree and Random Forest using the same cleaned dataset, four required input features, stratified train-test split, evaluation metrics, and untouched test set. Each classifier is placed within a scikit-learn Pipeline and tuned using five-fold StratifiedKFold GridSearchCV. This provides a consistent comparison and prevents the final test records from influencing model selection.")
    add_body(doc, "GridSearchCV evaluates different depth, minimum-leaf, and class-weight configurations, with Approved-class F1-score used as the primary model-selection metric. Both selected pipelines are then evaluated using accuracy, precision, recall, and F1-score and integrated into a Streamlit application. Because the dataset source does not fully document its institution, collection process, currency, or authenticity, the findings are limited to this dataset and are not presented as evidence of real-world banking performance.")


def add_methodology(doc):
    add_heading(doc, "3. Methodology", 1)
    add_heading(doc, "3.1 System Flowchart / Activity Diagram", 2)
    p = doc.add_paragraph("Figure 3.1. End-to-end loan approval prediction workflow.", style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(WORKFLOW_PATH), width=Inches(6.35))
    add_body(doc, "The workflow begins by loading and validating loan_approval_dataset.csv. Whitespace is removed from headers and category values, the identifier and labels are checked, and the four selected features are extracted. The data is then divided using a reproducible stratified 70:30 split. The test set is held aside while five-fold GridSearchCV tunes Decision Tree and Random Forest models using only the training data. Each selected model is evaluated once on the untouched test set, saved with Joblib, and loaded by the Streamlit application.")

    add_heading(doc, "3.2 Description and Analysis of Dataset", 2)
    add_heading(doc, "Dataset Source and Structure", 3)
    add_body(doc, "The final dataset is loan_approval_dataset.csv from the Kaggle Loan Approval Prediction Dataset page published by Archit Sharma. It contains 4,269 rows and 13 columns. The target column, loan_status, contains Approved and Rejected labels. The loan_id column uniquely identifies an application and is excluded from model training. The source page lists an MIT licence, but does not clearly document the original institution, collection process, monetary units, or whether the observations are authentic or generated.")
    add_table(
        doc,
        "Table 3.1. Original dataset schema and final modelling role.",
        ["Column", "Type", "Final role"],
        [
            ("loan_id", "Identifier", "Excluded; traceability only"),
            ("no_of_dependents", "Numeric", "Optional descriptive field"),
            ("education", "Categorical", "Optional descriptive field"),
            ("self_employed", "Categorical", "Optional descriptive field"),
            ("income_annum", "Numeric", "Required model input"),
            ("loan_amount", "Numeric", "Required model input"),
            ("loan_term", "Numeric", "Required model input"),
            ("cibil_score", "Numeric", "Required model input"),
            ("residential_assets_value", "Numeric", "Optional; excluded"),
            ("commercial_assets_value", "Numeric", "Optional; excluded"),
            ("luxury_assets_value", "Numeric", "Optional; excluded"),
            ("bank_asset_value", "Numeric", "Optional; excluded"),
            ("loan_status", "Categorical", "Prediction target"),
        ],
        [3300, 1800, 4260],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_heading(doc, "Dataset Characteristics", 3)
    add_body(doc, "The source file contains 2,656 Approved applications (62.2%) and 1,613 Rejected applications (37.8%). No missing values or exact duplicate rows were found. The moderate class difference does not require synthetic oversampling for this experiment; stratification preserves the distribution during splitting, while precision, recall, and F1-score complement accuracy.")
    add_table(
        doc,
        "Table 3.2. Descriptive statistics for the four required inputs.",
        ["Feature", "Minimum", "Maximum", "Mean"],
        [
            ("Annual income", "200,000", "9,900,000", "5,059,123.92"),
            ("Loan amount", "300,000", "39,500,000", "15,133,450.46"),
            ("Loan term", "2", "20", "10.90"),
            ("CIBIL score", "300", "900", "599.94"),
        ],
        [3150, 2070, 2070, 2070],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_callout(doc, "Unit limitation", "The Kaggle source does not state the monetary currency or the unit of loan_term. Values are therefore reported in dataset units and should not be relabelled as a specific currency or duration without source evidence.")

    add_heading(doc, "Data Preprocessing", 3)
    preprocessing = [
        "Load the CSV using whitespace-aware parsing and strip remaining whitespace from column names and category values.",
        "Validate the required identifier, target, labels, and numeric ranges.",
        "Confirm that loan_id values are unique and exclude the identifier from prediction.",
        "Detect duplicate applications and conflicting labels before the split.",
        "Select income_annum, loan_amount, loan_term, and cibil_score in a fixed column order.",
        "Exclude seven optional low-importance source features from the final model and Streamlit form.",
        "Create a stratified 70:30 train-test split with random_state=42.",
        "Verify that identical four-feature applications do not occur in both subsets.",
        "Save the processed features and labels as reusable Joblib artifacts.",
    ]
    add_numbered_list(doc, preprocessing)
    add_body(doc, "The original asset columns include 28 records where residential_assets_value equals -100000. The value is preserved in the original CSV and documented as an anomaly, but the affected asset field is excluded from the final model. No unsupported replacement value is introduced.")

    add_heading(doc, "Selected Inputs and Train-Test Split", 3)
    add_body(doc, "All four selected inputs are numerical. Tree-based models do not require feature scaling, so the shared ColumnTransformer passes the values through unchanged. Keeping preprocessing inside each scikit-learn Pipeline ensures that the same schema is applied during cross-validation, final training, saved-model inference, and Streamlit prediction.")
    add_table(
        doc,
        "Table 3.3. Stratified dataset distribution.",
        ["Dataset stage", "Approved", "Rejected", "Total"],
        [
            ("Validated dataset", "2,656", "1,613", "4,269"),
            ("Training set", "1,859", "1,129", "2,988"),
            ("Untouched test set", "797", "484", "1,281"),
        ],
        [3600, 1920, 1920, 1920],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_body(doc, "Preliminary analysis shows that CIBIL score has a very strong relationship with the target: approximately 10.36% of applications below 550 are Approved, compared with approximately 99.48% at 550 or above. This explains much of the model performance and is a major limitation when interpreting the results.")

    add_heading(doc, "3.3 Algorithm Selection and Description", 2)
    add_heading(doc, "Decision Tree Classifier", 3)
    add_body(doc, "A Decision Tree is a non-parametric supervised classifier that divides the feature space using a sequence of decision rules. Each internal node selects a feature and split point, and each terminal leaf produces a class prediction. Its main advantage is interpretability because the learned conditions can be visualized and explained. Its main weakness is instability and overfitting when the tree grows too deeply. The project therefore tunes maximum depth, minimum samples per leaf, and class weighting.")
    add_table(
        doc,
        "Table 3.4. Decision Tree hyperparameter search space.",
        ["Parameter", "Values evaluated"],
        [
            ("max_depth", "3, 4, 5, 6, 8, 10, None"),
            ("min_samples_leaf", "1, 2, 5, 10"),
            ("class_weight", "None, balanced"),
        ],
        [3100, 6260],
    )
    add_body(doc, "The search contains 56 parameter combinations. With five validation folds, GridSearchCV performs 280 fits before refitting the best configuration on the complete training set.")

    add_heading(doc, "Random Forest Classifier", 3)
    add_body(doc, "Random Forest combines predictions from many Decision Trees trained using random samples and randomized feature selection. Aggregating 300 trees reduces the sensitivity of a single tree to small changes in the training data and usually improves generalization. The trade-off is reduced interpretability and a larger saved model. The number of estimators is fixed at 300, while tree depth, minimum leaf size, and class weighting are tuned.")
    add_table(
        doc,
        "Table 3.5. Random Forest hyperparameter search space.",
        ["Parameter", "Values evaluated"],
        [
            ("n_estimators", "300 (fixed)"),
            ("max_depth", "None, 10, 20"),
            ("min_samples_leaf", "1, 2"),
            ("class_weight", "None, balanced"),
        ],
        [3100, 6260],
    )
    add_body(doc, "The search contains 12 parameter combinations and therefore performs 60 five-fold validation fits before refitting the selected configuration.")

    add_heading(doc, "Cross-Validation and Model Selection", 3)
    add_body(doc, "GridSearchCV uses five-fold StratifiedKFold with shuffle=True and random_state=42. In every iteration, four folds are used for training and one fold for validation. Stratification preserves the Approved and Rejected proportions. Accuracy, Approved-class precision, recall, and F1-score are recorded, while F1-score is the refit and model-selection metric. After selection, the best pipeline is fitted using the entire 2,988-row training set and evaluated once on the untouched 1,281-row test set.")

    add_heading(doc, "3.4 Evaluation Metrics", 2)
    add_body(doc, "Approved is treated as the positive class. TP represents an Approved application predicted Approved; TN represents a Rejected application predicted Rejected; FP represents a Rejected application predicted Approved; and FN represents an Approved application predicted Rejected.")
    add_table(
        doc,
        "Table 3.6. Evaluation metrics used in the project.",
        ["Metric", "Formula", "Interpretation"],
        [
            ("Accuracy", "(TP + TN) / (TP + TN + FP + FN)", "Overall proportion predicted correctly"),
            ("Precision", "TP / (TP + FP)", "Correctness of predicted approvals"),
            ("Recall", "TP / (TP + FN)", "Coverage of actual approvals"),
            ("F1-score", "2 x (Precision x Recall) / (Precision + Recall)", "Balance between precision and recall"),
        ],
        [1600, 3350, 4410],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body(doc, "Only these four metrics are retained in the final model-comparison result. F1-score is used for tuning because it combines precision and recall, while the other three values provide complementary views of overall correctness and Approved-class performance.")


def add_results(doc):
    add_heading(doc, "4. Results and Discussion", 1)
    add_heading(doc, "4.1 Results", 2)
    add_heading(doc, "Dataset Preparation Results", 3)
    add_body(doc, "All 4,269 source rows passed the final schema and range validation. No missing or duplicate records were removed. The preparation stage selected four numerical inputs, created 2,988 training rows and 1,281 test rows, and confirmed that identical four-feature applications did not cross the split boundary.")

    add_heading(doc, "Cross-Validation Results", 3)
    add_table(
        doc,
        "Table 4.1. Mean five-fold validation performance of the selected configurations.",
        ["Model", "Accuracy", "Precision", "Recall", "F1-score"],
        [
            ("Decision Tree", "98.26%", "98.60%", "98.60%", "98.60%"),
            ("Random Forest", "98.03%", "98.03%", "98.82%", "98.42%"),
        ],
        [2640, 1680, 1680, 1680, 1680],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_body(doc, "The Decision Tree achieved the slightly higher mean validation F1-score, so its selected configuration used class_weight=balanced, max_depth=10, and min_samples_leaf=5. The selected Random Forest used class_weight=None, unrestricted depth, and min_samples_leaf=1. The validation scores were close, so final comparison was based on the common untouched test set.")

    add_heading(doc, "Final Test-Set Results", 3)
    add_table(
        doc,
        "Table 4.2. Final performance on the untouched 1,281-row test set.",
        ["Model", "Accuracy", "Precision", "Recall", "F1-score"],
        [
            ("Random Forest", "98.75%", "98.51%", "99.50%", "99.00%"),
            ("Decision Tree", "98.28%", "98.14%", "99.12%", "98.63%"),
        ],
        [2640, 1680, 1680, 1680, 1680],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    p = doc.add_paragraph("Figure 4.1. Decision Tree and Random Forest test-set performance.", style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ROOT / "models" / "comparison_chart.png"), width=Inches(5.9))

    add_heading(doc, "4.2 Discussion / Interpretation", 2)
    add_body(doc, "Random Forest achieved the strongest final result for every retained metric. Its accuracy of 98.75% exceeded the Decision Tree by 0.47 percentage points. It also achieved 98.51% precision, 99.50% recall, and a 99.00% F1-score. Random Forest is therefore selected as the primary model for this dataset.")
    add_body(doc, "Decision Tree nevertheless performed strongly, with 98.28% accuracy and a 98.63% F1-score. Its simpler structure makes it valuable as an interpretable comparison model. The selected depth limit of 10 and minimum leaf size of five reduce the risk of an excessively complex tree, while balanced class weighting adjusts the contribution of the two target classes during training.")
    add_body(doc, "The small gap between the models should be interpreted with caution. Random Forest's higher test performance is consistent with the ensemble principle described by Breiman (2001), because combining many randomized trees can reduce the instability of an individual tree. However, a difference observed on one 1,281-row test set does not prove that Random Forest will always outperform Decision Tree on future or externally collected applications.")
    add_body(doc, "The most important limitation is the relationship between CIBIL score and loan_status. Approval is approximately 10.36% below a CIBIL score of 550 and 99.48% at 550 or above. The models can therefore achieve very high performance by learning a near-threshold pattern that may have been built into the dataset. These scores demonstrate successful pattern reproduction within the supplied records; they do not establish realistic credit-risk assessment or fairness across real applicants.")
    add_body(doc, "The Streamlit prototype presents both model predictions and indicates whether the models agree. The form intentionally requires only annual income, loan amount, loan term, and CIBIL score. Optional descriptive and asset fields are omitted because the trained pipelines do not use them. This keeps the interface consistent with the saved model schema and prevents users from assuming that unused information influences a prediction.")


def add_conclusion(doc):
    add_heading(doc, "5. Conclusion", 1)
    add_heading(doc, "5.1 Achievements", 2)
    add_body(doc, "The project successfully developed a complete supervised machine-learning system for predicting whether a loan application is Approved or Rejected. It implemented dataset validation, whitespace normalization, feature selection, a reproducible stratified split, leakage-safe pipelines, five-fold GridSearchCV, model training, four-metric evaluation, artifact generation, automated testing, and Streamlit-based inference.")
    add_body(doc, "Both optimized classifiers performed strongly on the untouched test set. Random Forest achieved 98.75% accuracy, 98.51% precision, 99.50% recall, and 99.00% F1-score. Decision Tree achieved 98.28% accuracy, 98.14% precision, 99.12% recall, and 98.63% F1-score. Random Forest was therefore selected as the primary model because it obtained the highest result for all four evaluation metrics, while Decision Tree remained a suitable interpretable comparison model.")
    add_body(doc, "The project fulfilled its objectives by developing two supervised classification solutions, optimizing and evaluating them under the same controlled conditions, comparing their predictive performance, saving the fitted models, and integrating both pipelines into a working four-input prediction application.")

    add_heading(doc, "5.2 Limitations and Future Work", 2)
    limitations = [
        "Dataset provenance is incomplete. The originating institution, collection process, currency, term unit, and authenticity of the observations are not clearly documented.",
        "The target describes historical approval decisions rather than later repayment or default. The system therefore does not measure credit risk directly.",
        "CIBIL score almost determines the label, so the very high results may reflect a simple dataset rule rather than a realistic lending process.",
        "Evaluation uses one dataset and one held-out test set. No external institution or later time period is available for validation.",
        "The reduced four-input model improves usability but omits information that could matter in a real lending context. Conversely, adding demographic variables could introduce fairness and discrimination concerns.",
        "Random Forest is less interpretable than a single Decision Tree, and the interface does not yet provide local explanations for individual predictions.",
    ]
    for item in limitations:
        add_bullet(doc, item)
    add_body(doc, "Future work should evaluate a well-documented real or carefully generated dataset, distinguish approval prediction from default-risk prediction, test performance across external data, examine fairness across legally and ethically relevant groups, add calibrated probability estimates, and implement model explanations such as permutation importance or SHAP. Threshold selection should be tied to an explicitly defined cost of false approvals and false rejections. Any operational system would also require human review, governance, audit trails, security controls, and compliance with applicable lending and data-protection rules.")


def add_references(doc):
    add_heading(doc, "References and Sources", 1)
    refs = [
        "Bhakti, D. S., Prasetyo, A., & Arsi, P. (2024). Implementation of hyperparameter tuning in Random Forest algorithm for loan approval prediction. Jurnal Teknik Informatika, 5(4), 63-69. https://doi.org/10.52436/1.jutif.2024.5.4.2032",
        "Breiman, L. (2001). Random forests. Machine Learning, 45, 5-32. https://doi.org/10.1023/A:1010933404324",
        "Lakshmi, S. R., Gajendra, S., & Mohanraju, V. S. (2025). An enhanced ensemble-based framework for loan approval prediction using machine learning. International Journal of Environmental Sciences, 3164-3171. https://doi.org/10.64252/cnrb6282",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830. https://jmlr.org/papers/v12/pedregosa11a.html",
        "Scikit-learn developers. (2026). Decision Trees. Scikit-learn 1.9.0 documentation. https://scikit-learn.org/stable/modules/tree.html",
        "Scikit-learn developers. (2026). Ensembles: Random forests. Scikit-learn 1.9.0 documentation. https://scikit-learn.org/stable/modules/ensemble.html",
        "Sinap, V. (2024). A comparative study of loan approval prediction using machine learning methods. Gazi University Journal of Science Part C: Design and Technology, 12(2), 644-663. https://doi.org/10.29109/gujsc.1455978",
        "Streamlit. (2026). Streamlit documentation. https://docs.streamlit.io/",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.10
        set_run_font(p.add_run(ref), size=10)

    add_heading(doc, "Dataset Source", 2)
    add_body(doc, "Sharma, A. (n.d.). Loan Approval Prediction Dataset [Data set]. Kaggle. https://www.kaggle.com/datasets/architsharma01/loan-approval-prediction-dataset")
    add_body(doc, "The submitted CSV contains 4,269 labelled applications. Its SHA-256 hash is 4B5CD093D178378F4CFA8C107ADB6E599B88BE9D8A3B51F3B99C0D5914154E54. The source page lists an MIT licence. Provenance and quality limitations are documented in data/README.md.")

    add_heading(doc, "Development Tools", 2)
    add_body(doc, "The project was developed with Python, pandas, scikit-learn, Joblib, Matplotlib, Seaborn, and Streamlit. The dependency ranges are recorded in requirements.txt. Model summaries record scikit-learn version 1.9.0 for the successful training run.")

    add_heading(doc, "Source Code", 2)
    add_body(doc, "Repository: https://github.com/cheeweng20/AI-spam-classifier-project")
    add_body(doc, "The repository name is a legacy name from the earlier spam-classification title; its current content implements loan approval prediction. The main modules are src/settings.py, src/prepare_data.py, src/training_utils.py, src/train_decision_tree.py, src/train_random_forest.py, src/compare_models.py, streamlit_app.py, and tests/test_data_pipeline.py.")


def add_appendix(doc):
    add_heading(doc, "Appendix A: Project File Guide", 1)
    add_table(
        doc,
        "Table A.1. Main project files and responsibilities.",
        ["File", "Purpose"],
        [
            ("data/loan_approval_dataset.csv", "Final raw dataset"),
            ("data/README.md", "Dataset schema, provenance, quality notes, and limitations"),
            ("src/settings.py", "Paths, random seed, labels, selected features, and excluded source fields"),
            ("src/prepare_data.py", "Cleaning, validation, feature selection, split, and processed artifacts"),
            ("src/training_utils.py", "Shared pipelines, GridSearchCV, four metrics, plots, and persistence"),
            ("src/train_decision_tree.py", "Decision Tree search and training"),
            ("src/train_random_forest.py", "Random Forest search and training"),
            ("src/compare_models.py", "Four-metric comparison table and chart"),
            ("streamlit_app.py", "Four-input prediction interface"),
            ("tests/test_data_pipeline.py", "Automated validation and inference tests"),
        ],
        [3600, 5760],
    )
    add_heading(doc, "Reproduction Commands", 2)
    commands = [
        "python -m pip install -r requirements.txt",
        "python src/prepare_data.py",
        "python src/train_decision_tree.py",
        "python src/train_random_forest.py",
        "python src/compare_models.py",
        "python -m unittest discover -s tests -v",
        "python -m streamlit run streamlit_app.py",
    ]
    for command in commands:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(command)
        set_run_font(run, name="Consolas", size=9.5, color=NAVY)


def set_core_properties(doc):
    props = doc.core_properties
    props.title = "Loan Approval Prediction Assignment Documentation"
    props.subject = "Comparison of Decision Tree and Random Forest classifiers"
    props.author = "Ee Wan Yin and Lam Chee Weng"
    props.keywords = "loan approval, supervised machine learning, decision tree, random forest"
    props.comments = "Generated as an editable assignment draft; student review and verification required."


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_workflow()
    doc = Document()
    configure_document(doc)
    set_core_properties(doc)
    add_cover(doc)
    add_intro(doc)
    add_related_work(doc)
    add_methodology(doc)
    add_results(doc)
    add_conclusion(doc)
    add_references(doc)
    add_appendix(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
